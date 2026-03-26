import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_doc():
    doc = Document()
    
    # -----------------------------------------------------------------------
    # Title Page
    # -----------------------------------------------------------------------
    title = doc.add_heading('Automated Enterprise Loan Approval System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Comprehensive Project Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # -----------------------------------------------------------------------
    # Content Source
    # -----------------------------------------------------------------------
    doc_dir = "documentation"
    files = [
        "01_Beginner_Overview.md",
        "02_Technical_Architecture.md",
        "03_Technologies_Used.md",
        "04_User_Guide.md"
    ]
    
    for filename in files:
        filepath = os.path.join(doc_dir, filename)
        if not os.path.exists(filepath):
            continue
            
        with open(filepath, "r", encoding="utf-8") as f:
            lines = f.readlines()
            
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("* ") or line.startswith("- "):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith("1. ") or line.startswith("2. "):
                doc.add_paragraph(line[3:], style='List Number')
            else:
                doc.add_paragraph(line)
        
        doc.add_page_break()
    
    # -----------------------------------------------------------------------
    # Save
    # -----------------------------------------------------------------------
    output_path = os.path.join(doc_dir, "Project_Documentation.docx")
    doc.save(output_path)
    print(f"✅ Word document created successfully at: {output_path}")

if __name__ == "__main__":
    create_word_doc()
